import datetime
import io
import json
import os
import re
import base64
from collections import OrderedDict

from flask import Flask, jsonify, render_template, request, send_file

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError:  # pragma: no cover
    Document = None

try:  # pragma: no cover
    import pdfplumber
except ImportError:  # pragma: no cover
    pdfplumber = None

try:  # pragma: no cover
    from PyPDF2 import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app = Flask(__name__)


def slugify(value):
    cleaned = "".join(ch.lower() if ch.isalnum() else "-" for ch in value).strip("-")
    return "-".join(part for part in cleaned.split("-") if part) or "item"


SEED_LIBRARY = {
    "company": {
        "name": "RELIANCE ENGINEERING",
        "quote_title": "QUOTE",
        "tagline": "Interactive Techno Solution Pvt. Ltd.",
        "contact_person": "Muhammad Ali",
        "phone": "03228233602",
        "email": "contact@intechnosol.com",
        "address": "Plot 11C, Lane 1, Sehar Commercial, DHA Phase 6, Karachi",
        "footer_line_1": "Plot 11C, Lane 1, Sehar Commercial, DHA Phase 6, Karachi",
        "footer_line_2": "T: +92 353 4844 545   E: contact@intechnosol.com   W: intechnosol.com",
        "terms": [
            "Interactive Techno Solutions Pvt. Ltd.",
            "Prices are exclusive of all taxes.",
            "100% Advance Payment.",
        ],
        "closing": [
            "Thank you and assuring you of our best services at all time remain.",
            "Yours faithfully",
            "Interactive Techno Solutions Pvt. Ltd.",
        ],
    },
    "specOptions": [
        "Nos", "Set", "Job", "Kg", "1kg", "2kg", "5kg", "L", "1L", "2L", "5L",
        "500ml", "3/4", "5mm", "6mm", "10mm", "12mm", "16mm", "25mm",
        "30A", "32A", "63A", "100A", "40x40", "60x60", "75mm",
        '1.5"x3"', '2"x2"', '2"x5"', '2.5"x5"', '3.5"', '4.5"',
        "5x5", "6x12x18", "10x10", "14 gauge", "16 gauge", "24V", "48V",
        "1000V DC", "IP65", "2.5kW", "5.1kWh", "6kW", "10kW", "12000W",
        "570W", "585W", "590W", "605W", "615W", "645W", "720W",
        "94x46", "96x45", "2500 PSI",
    ],
    "categories": [
        {
            "id": "elevated-structure",
            "name": "ELEVATED STRUCTURE",
            "color": "#8d6a2f",
            "sortOrder": 1,
            "subcategories": [
                {
                    "id": "fabricated-structure-steel",
                    "name": "Fabricated Structure Steel",
                    "items": [
                        {"id": "pole-4-5-dia", "name": 'Pole 4.5" Dia', "description": 'Pole 4.5" dia in 2.5mm', "defaultSpec": '4.5"', "defaultUnitPrice": ""},
                        {"id": "base-girder-2-5x5", "name": 'Base Girder 2.5"x5"', "description": 'Base girder 2.5"x5" JSI', "defaultSpec": '2.5"x5"', "defaultUnitPrice": ""},
                        {"id": "plate-placing-girder-2-5x5", "name": 'Plate Placing Girder 2.5"x5"', "description": 'Plate placing girder 2.5"x5" JSI', "defaultSpec": '2.5"x5"', "defaultUnitPrice": ""},
                        {"id": "top-girder-2-5x5", "name": 'Top Girder 2.5"x5"', "description": 'Top girder 2.5"x5" section', "defaultSpec": '2.5"x5"', "defaultUnitPrice": ""},
                        {"id": "side-support-angle-2x2", "name": 'Cantilever Angle 2"x2"', "description": 'Structure side support and cantilever angle 2"x2" in 4.5mm', "defaultSpec": '2"x2"', "defaultUnitPrice": ""},
                        {"id": "channel-2x5-16g", "name": 'Channel 2"x5"', "description": 'Channel for plate placing 2"x5" in 16 gauge', "defaultSpec": '2"x5"', "defaultUnitPrice": ""},
                    ],
                },
                {
                    "id": "plates-anchors-fasteners",
                    "name": "Plates, Anchors & Fasteners",
                    "items": [
                        {"id": "base-plate-10x10", "name": 'Base Plate 10"x10"', "description": 'Base plate 10"x10" in 10mm', "defaultSpec": "10x10", "defaultUnitPrice": ""},
                        {"id": "top-plate-5x5", "name": 'Top Plate 5"x5"', "description": 'Top plate 5"x5" in 6mm', "defaultSpec": "5x5", "defaultUnitPrice": ""},
                        {"id": "rawal-bolt-10x75", "name": "Rawal Bolt 10mm x 75mm", "description": "Rawal bolt 10mm x 75mm", "defaultSpec": "75mm", "defaultUnitPrice": ""},
                        {"id": "anchor-bolt-4in-12mm", "name": 'Anchor Bolt 4" in 12mm', "description": 'Anchor bolt 4" in 12mm', "defaultSpec": "12mm", "defaultUnitPrice": ""},
                        {"id": "nut-bolt-ss316", "name": "Nut Bolt SS316", "description": "Nut bolt for solar SS316 No. 8", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "mid-clamp", "name": "Mid Clamp", "description": "Mid clamp for panel mounting", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "end-clamp", "name": "End Clamp", "description": "End clamp for panel mounting", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "u-clamp", "name": "U Clamp", "description": "U clamp for structure mounting", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "aluminum-tz-section", "name": "Aluminum TZ Section", "description": "Aluminum TZ section for plate clamp", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                    ],
                },
                {
                    "id": "finishing-materials",
                    "name": "Finishing Materials",
                    "items": [
                        {"id": "checker-sheet", "name": "Checker Sheet", "description": "Checker sheet for fabricated structure", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "epoxy-red-oxide", "name": "Buxly Epoxy Red Oxide", "description": "Buxly epoxy red oxide", "defaultSpec": "500ml", "defaultUnitPrice": ""},
                        {"id": "epoxy-smoke-grey", "name": "Smoke Grey Finish", "description": "Finishing with epoxy smoke grey (Nippon)", "defaultSpec": "500ml", "defaultUnitPrice": ""},
                        {"id": "paint-brush", "name": "Paint Brush", "description": "Paint brush for primer and finish", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                    ],
                },
                {
                    "id": "mechanical-civil",
                    "name": "Mechanical Installation & Civil",
                    "items": [
                        {"id": "mechanical-installation-labor", "name": "Mechanical Structure Installation Charges", "description": "Mechanical structure installation labor charges", "defaultSpec": "Job", "defaultUnitPrice": ""},
                        {"id": "civil-work", "name": "Civil Work", "description": "Civil work related to structure installation", "defaultSpec": "Job", "defaultUnitPrice": ""},
                        {"id": "curve-stone-6x12x18", "name": "Concrete / Curve Stone", "description": "Concrete or curve stone 6x12x18, 2500 PSI", "defaultSpec": "6x12x18", "defaultUnitPrice": ""},
                        {"id": "lightning-earthing-structure", "name": "Structure Earthing", "description": "Earthing for structure and lightning arrestor", "defaultSpec": "Job", "defaultUnitPrice": ""},
                    ],
                },
            ],
        },
        {
            "id": "solar-system",
            "name": "SOLAR SYSTEM",
            "color": "#1f6f5f",
            "sortOrder": 2,
            "subcategories": [
                {
                    "id": "solar-panels",
                    "name": "Solar Panels",
                    "items": [
                        {"id": "canadian-solar-tier1", "name": "Canadian Solar Tier 1", "description": "Canadian Solar Tier 1 A Grade panel", "defaultSpec": "605W", "defaultUnitPrice": ""},
                        {"id": "jinko-solar-tier1", "name": "Jinko Solar Tier 1", "description": "Jinko Solar Tier 1 A Grade panel", "defaultSpec": "605W", "defaultUnitPrice": ""},
                        {"id": "longi-solar-tier1", "name": "Longi Solar Tier 1", "description": "Longi Solar Tier 1 A Grade panel", "defaultSpec": "605W", "defaultUnitPrice": ""},
                        {"id": "ja-solar-tier1", "name": "JA Solar Tier 1", "description": "JA Solar Tier 1 A Grade panel", "defaultSpec": "605W", "defaultUnitPrice": ""},
                        {"id": "trina-solar-tier1", "name": "Trina Solar Tier 1", "description": "Trina Solar Tier 1 A Grade panel", "defaultSpec": "605W", "defaultUnitPrice": ""},
                        {"id": "mono-tier1-570", "name": "Mono Tier 1 570W", "description": "Mono Tier 1 A Grade Canadian/Jinko/Longi, TUV certified, ISO 9001, EN 61215, EN 61730", "defaultSpec": "570W", "defaultUnitPrice": ""},
                        {"id": "mono-tier1-585", "name": "Mono Tier 1 585W", "description": "Mono solar panel 585W", "defaultSpec": "585W", "defaultUnitPrice": ""},
                        {"id": "mono-tier1-590", "name": "Mono Tier 1 590W", "description": "Mono solar panel 590W", "defaultSpec": "590W", "defaultUnitPrice": ""},
                        {"id": "mono-tier1-605", "name": "Mono Tier 1 605W", "description": "Mono Tier 1 A Grade 605W module", "defaultSpec": "605W", "defaultUnitPrice": ""},
                        {"id": "mono-tier1-615", "name": "Mono Tier 1 615W", "description": "Mono solar panel 615W", "defaultSpec": "615W", "defaultUnitPrice": ""},
                        {"id": "mono-tier1-645", "name": "Mono Tier 1 645W", "description": "Mono solar panel 645W", "defaultSpec": "645W", "defaultUnitPrice": ""},
                        {"id": "bifacial-605", "name": "Bifacial 605W", "description": "605W bifacial solar panel", "defaultSpec": "605W", "defaultUnitPrice": ""},
                        {"id": "mono-tier1-720", "name": "Mono Tier 1 720W", "description": "720W high output solar panel", "defaultSpec": "720W", "defaultUnitPrice": ""},
                        {"id": "panel-size-96x45", "name": 'Panel Size 96" x 45"', "description": 'Panel dimensional note 96" x 45"', "defaultSpec": "96x45", "defaultUnitPrice": ""},
                        {"id": "panel-size-94x46", "name": 'Panel Size 94" x 46"', "description": 'Panel dimensional note 94" x 46"', "defaultSpec": "94x46", "defaultUnitPrice": ""},
                    ],
                },
                {
                    "id": "inverters",
                    "name": "Inverters",
                    "items": [
                        {"id": "sungrow-6kw", "name": "Sungrow 6kW Hybrid Inverter", "description": "6kW IP65 inverter, max PV 12000W, 5-year warranty", "defaultSpec": "6kW", "defaultUnitPrice": ""},
                        {"id": "huawei-inverter", "name": "Huawei Inverter", "description": "Huawei inverter option", "defaultSpec": "6kW", "defaultUnitPrice": ""},
                        {"id": "on-grid-inverter", "name": "On-Grid Inverter", "description": "On-grid inverter option", "defaultSpec": "6kW", "defaultUnitPrice": ""},
                        {"id": "off-grid-inverter", "name": "Off-Grid Inverter", "description": "Off-grid inverter option", "defaultSpec": "6kW", "defaultUnitPrice": ""},
                        {"id": "hybrid-inverter", "name": "Hybrid Inverter", "description": "Hybrid inverter option, IP65 rated", "defaultSpec": "IP65", "defaultUnitPrice": ""},
                        {"id": "solis-inverter", "name": "Solis Inverter", "description": "Solis inverter option", "defaultSpec": "6kW", "defaultUnitPrice": ""},
                        {"id": "growatt-inverter", "name": "Growatt Inverter", "description": "Growatt inverter option", "defaultSpec": "6kW", "defaultUnitPrice": ""},
                        {"id": "inverex-inverter", "name": "Inverex Inverter", "description": "Inverex inverter option", "defaultSpec": "6kW", "defaultUnitPrice": ""},
                    ],
                },
                {
                    "id": "batteries",
                    "name": "Batteries",
                    "items": [
                        {"id": "lithium-5-1", "name": "Lithium 5.1kWh 48V", "description": "Lithium battery 5.1kWh, 48V, 7-year warranty (Blanty/Dyness)", "defaultSpec": "5.1kWh", "defaultUnitPrice": ""},
                        {"id": "lead-acid-battery", "name": "Lead Acid Battery", "description": "Lead acid battery option", "defaultSpec": "24V", "defaultUnitPrice": ""},
                        {"id": "dyness-battery", "name": "Dyness Battery", "description": "Dyness lithium battery option", "defaultSpec": "48V", "defaultUnitPrice": ""},
                        {"id": "blanty-battery", "name": "Blanty Battery", "description": "Blanty lithium battery option", "defaultSpec": "48V", "defaultUnitPrice": ""},
                        {"id": "battery-wall-bracket", "name": "Battery Wall Bracket", "description": "Battery wall bracket / support", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "dc-battery-wire-25mm", "name": "DC Battery Wire", "description": "DC 25mm battery wire", "defaultSpec": "25mm", "defaultUnitPrice": ""},
                    ],
                },
                {
                    "id": "electrical-wiring",
                    "name": "Electrical Wiring",
                    "items": [
                        {"id": "solar-wire-6mm", "name": "Solar Wire 6mm XLPE / XLPO", "description": "Wire 6mm XLPE / XLPO tin coated for solar panel wiring", "defaultSpec": "6mm", "defaultUnitPrice": ""},
                        {"id": "ac-wire-6mm", "name": "AC Wire 6mm", "description": "AC wire for main panel and power connectivity", "defaultSpec": "6mm", "defaultUnitPrice": ""},
                        {"id": "ac-wire-10mm", "name": "AC Wire 10mm", "description": "AC wire 10mm for heavier power connectivity", "defaultSpec": "10mm", "defaultUnitPrice": ""},
                        {"id": "lug-10-16-25", "name": "Cable Lug", "description": "Cable lug 10 / 16 / 25 mm", "defaultSpec": "25mm", "defaultUnitPrice": ""},
                    ],
                },
                {
                    "id": "protection-switching",
                    "name": "Protection & Switching",
                    "items": [
                        {"id": "main-breaker-ac-input", "name": "Main Breaker AC Input", "description": "Main breaker for AC input", "defaultSpec": "63A", "defaultUnitPrice": ""},
                        {"id": "main-breaker-ac-output", "name": "Main Breaker AC Output", "description": "Main breaker for AC output", "defaultSpec": "63A", "defaultUnitPrice": ""},
                        {"id": "dc-breaker-pv", "name": "DC Breaker for PV", "description": "DC breaker for solar panel / PV input", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "distribution-box-ip65", "name": "Distribution Box IP65", "description": "Distribution box IP65 branded", "defaultSpec": "IP65", "defaultUnitPrice": ""},
                        {"id": "changeover-selector-63a", "name": "Change Over Selector", "description": "Change over selector for power bypass", "defaultSpec": "63A", "defaultUnitPrice": ""},
                        {"id": "mc4-connector", "name": "MC4 Connector", "description": "MC4 connector 1000V DC", "defaultSpec": "1000V DC", "defaultUnitPrice": ""},
                        {"id": "changeover-opas", "name": "Changeover (Opas)", "description": "Changeover switch (Opas)", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                    ],
                },
                {
                    "id": "fittings-accessories",
                    "name": "Fittings & Accessories",
                    "items": [
                        {"id": "pvc-pipe-3-4", "name": 'PVC Pipe 3/4"', "description": 'PVC fitting pipe 3/4" with band, socket, and saddle', "defaultSpec": "3/4", "defaultUnitPrice": ""},
                        {"id": "socket", "name": "Socket", "description": "Socket fitting for electrical work", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "screw-gitti", "name": "Screw / Gitti", "description": "Screw and gitti accessories", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "heat-sleeve", "name": "Heat Sleeve", "description": "Heat sleeve / insulation sleeve", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "duct-40x40", "name": "Duct 40x40", "description": "PVC duct / slotted duct 40x40", "defaultSpec": "40x40", "defaultUnitPrice": ""},
                        {"id": "duct-60x60", "name": "Duct 60x60", "description": "PVC duct / slotted duct 60x60", "defaultSpec": "60x60", "defaultUnitPrice": ""},
                        {"id": "flexible-pipe", "name": "Flexible Pipe", "description": "Flexible pipe for wire routing", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "cable-tie", "name": "Cable Tie", "description": "Cable tie for wiring management", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "lux-pvc-tape", "name": "PVC Tape", "description": "Lux PVC tape", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                        {"id": "insulation-tape", "name": "Insulation Tape", "description": "Insulation / tagging tape", "defaultSpec": "Nos", "defaultUnitPrice": ""},
                    ],
                },
                {
                    "id": "earthing-transport-installation",
                    "name": "Earthing, Transportation & Installation",
                    "items": [
                        {"id": "earthling", "name": "Earthing for Lightning Arrestor and Structure", "description": "Earthing for lightning arrestor and structure", "defaultSpec": "Job", "defaultUnitPrice": ""},
                        {"id": "transportation", "name": "Transportation", "description": "Transportation charges", "defaultSpec": "Job", "defaultUnitPrice": ""},
                        {"id": "installation", "name": "Installation", "description": "Installation charges", "defaultSpec": "Job", "defaultUnitPrice": ""},
                    ],
                },
            ],
        },
    ],
    "defaultQuote": {
        "preparedBy": "Muhammad Ali",
        "preparedPhone": "03228233602",
        "preparedEmail": "contact@intechnosol.com",
        "to": "",
        "proposalFor": "",
        "city": "",
        "quoteDate": "",
    },
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_table_borders(table, color="B8BFC6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def apply_cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"


def money(value):
    try:
        return f"{float(value):,.2f}"
    except (TypeError, ValueError):
        return "0.00"


def decode_data_url_image(data_url):
    if not data_url or not isinstance(data_url, str):
        return None
    if "," not in data_url:
        return None
    header, encoded = data_url.split(",", 1)
    if ";base64" not in header:
        return None
    try:
        return io.BytesIO(base64.b64decode(encoded))
    except Exception:
        return None


def parse_amount(value):
    cleaned = re.sub(r"[^0-9.\-]", "", str(value or ""))
    try:
        return float(cleaned) if cleaned else 0.0
    except ValueError:
        return 0.0


def normalize_text(value):
    return " ".join(str(value or "").replace("\n", " ").split()).strip()


def normalize_key(value):
    return normalize_text(value).lower()


def remove_prefix_case_insensitive(text, prefix):
    if normalize_key(text).startswith(normalize_key(prefix)):
        return text[len(prefix):].strip()
    return text


def extract_quote_meta_from_intro(left_text, right_text):
    lines = [line.strip() for line in str(left_text or "").splitlines() if line.strip()]
    meta = {
        "preparedBy": "",
        "preparedPhone": "",
        "preparedEmail": "",
        "to": "",
        "proposalFor": "",
        "city": "",
        "quoteDate": normalize_text(right_text),
    }

    if len(lines) >= 4:
        meta["preparedBy"] = lines[1]
        meta["preparedPhone"] = lines[2]
        meta["preparedEmail"] = lines[3]

    for index, line in enumerate(lines):
        lower = line.lower()
        if lower.startswith("to"):
            meta["to"] = line.replace("TO", "", 1).strip(" :-")
            if index + 1 < len(lines):
                meta["proposalFor"] = lines[index + 1]
            if index + 2 < len(lines):
                meta["city"] = lines[index + 2]
            break

    return meta


def parse_main_table_rows(rows):
    lines = []
    current_category = ""

    for row in rows[1:]:
        values = [normalize_text(cell.text) for cell in row.cells]
        if not any(values):
            continue

        unique_values = {value for value in values if value}
        first_value = values[0]

        if len(unique_values) == 1 and first_value and "subtotal" not in first_value.lower():
            current_category = first_value
            continue

        if "subtotal" in first_value.lower():
            continue

        if re.fullmatch(r"\d+", first_value):
            quantity = parse_amount(values[5])
            unit_price = parse_amount(values[6])
            total = parse_amount(values[7]) or (quantity * unit_price)
            lines.append(
                {
                    "serialNo": int(first_value),
                    "category": values[1] or current_category or "GENERAL",
                    "subcategory": values[2] or "General Items",
                    "description": values[3],
                    "spec": values[4],
                    "quantity": quantity or 1,
                    "unitPrice": unit_price,
                    "lineTotal": total,
                }
            )

    return lines


def parse_summary_table(table):
    summary = {"subtotal": 0.0, "discountAmount": 0.0, "taxAmount": 0.0, "grandTotal": 0.0}
    labels = {
        "subtotal": "subtotal",
        "discount": "discountAmount",
        "tax": "taxAmount",
        "grand total": "grandTotal",
    }
    for row in table.rows:
        key = normalize_text(row.cells[0].text).lower()
        for label, summary_key in labels.items():
            if key == label:
                summary[summary_key] = parse_amount(row.cells[1].text)
    return summary


def dedupe_import_lines(lines):
    deduped = []
    seen = set()
    for line in lines or []:
        key = (
            line.get("serialNo"),
            normalize_key(line.get("category")),
            normalize_key(line.get("subcategory")),
            normalize_key(line.get("description")),
            normalize_key(line.get("spec")),
            round(parse_amount(line.get("quantity")), 4),
            round(parse_amount(line.get("unitPrice")), 4),
            round(parse_amount(line.get("lineTotal")), 4),
        )
        if key in seen:
            continue
        seen.add(key)
        deduped.append(line)
    return deduped


def parse_docx_import(file_stream):
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX import.")

    document = Document(file_stream)
    main_table = None
    summary_table = None
    quote = {}

    for table in document.tables:
        if not table.rows:
            continue
        header = [normalize_text(cell.text) for cell in table.rows[0].cells]
        if header[:3] == ["S/N", "Category", "Subcategory"]:
            main_table = table
        elif len(table.rows[0].cells) == 2 and normalize_text(table.rows[0].cells[0].text).lower() == "subtotal":
            summary_table = table
        elif len(table.rows[0].cells) == 2 and not quote:
            quote = extract_quote_meta_from_intro(table.rows[0].cells[0].text, table.rows[0].cells[1].text)

    if main_table is None:
        raise ValueError("No editable quotation table was found in this Word file.")

    return {
        "importedFrom": "DOCX",
        "quote": quote,
        "lines": dedupe_import_lines(parse_main_table_rows(main_table.rows)),
        "summary": parse_summary_table(summary_table) if summary_table else {},
    }


def build_import_library(library_payload=None):
    base = {
        "categories": SEED_LIBRARY["categories"],
        "specOptions": SEED_LIBRARY["specOptions"],
    }
    if not library_payload:
        return base
    return {
        "categories": library_payload.get("categories", base["categories"]),
        "specOptions": library_payload.get("specOptions", base["specOptions"]),
    }


def parse_pdf_text_rows(text_lines, library_payload=None):
    library = build_import_library(library_payload)
    category_names = [category["name"] for category in library["categories"]]
    category_map = {normalize_key(name): name for name in category_names}
    category_names_sorted = sorted(category_names, key=len, reverse=True)

    subcategories_by_category = {}
    all_subcategories = []
    for category in library["categories"]:
        names = [subcategory["name"] for subcategory in category.get("subcategories", [])]
        subcategories_by_category[normalize_key(category["name"])] = names
        all_subcategories.extend(names)
    all_subcategories_sorted = sorted(set(all_subcategories), key=len, reverse=True)

    spec_options = sorted(set(library.get("specOptions", [])), key=len, reverse=True)
    quote = {
        "preparedBy": "",
        "preparedPhone": "",
        "preparedEmail": "",
        "to": "",
        "proposalFor": "",
        "city": "",
        "quoteDate": "",
    }

    rows = []
    started = False
    current_category = ""
    pending = ""
    intro_lines = []

    def finalize_pending(buffer_text):
        nonlocal current_category
        buffer = normalize_text(buffer_text)
        if not buffer:
            return None
        if normalize_key(buffer) in {normalize_key(name) for name in category_names_sorted}:
            current_category = category_map[normalize_key(buffer)]
            return None
        if "subtotal" in normalize_key(buffer):
            return None

        amount_match = re.search(r"\s([0-9]+(?:\.[0-9]+)?)\s+([0-9,]+(?:\.[0-9]+)?)\s+([0-9,]+(?:\.[0-9]+)?)\s*$", buffer)
        if not amount_match:
            return None

        prefix = buffer[:amount_match.start()].strip()
        quantity = parse_amount(amount_match.group(1))
        unit_price = parse_amount(amount_match.group(2))
        total = parse_amount(amount_match.group(3)) or (quantity * unit_price)

        serial_match = re.match(r"^(\d+)\s+(.*)$", prefix)
        if not serial_match:
            return None

        serial_no = int(serial_match.group(1))
        remainder = serial_match.group(2).strip()

        spec = ""
        for option in spec_options:
            option_key = normalize_key(option)
            if normalize_key(remainder).endswith(option_key):
                cutoff = len(remainder) - len(option)
                if cutoff <= 0 or remainder[cutoff - 1].isspace():
                    remainder = remainder[:cutoff].strip()
                    spec = option
                    break

        category_name = current_category or "GENERAL"
        for name in category_names_sorted:
            if normalize_key(remainder).startswith(normalize_key(name)):
                category_name = category_map[normalize_key(name)]
                remainder = remove_prefix_case_insensitive(remainder, name)
                break

        subcategory_name = "General Items"
        category_subs = subcategories_by_category.get(normalize_key(category_name), [])
        for name in sorted(category_subs, key=len, reverse=True):
            if normalize_key(remainder).startswith(normalize_key(name)):
                subcategory_name = name
                remainder = remove_prefix_case_insensitive(remainder, name)
                break
        else:
            for name in all_subcategories_sorted:
                if normalize_key(remainder).startswith(normalize_key(name)):
                    subcategory_name = name
                    remainder = remove_prefix_case_insensitive(remainder, name)
                    break

        description = remainder.strip()
        if not description:
            description = subcategory_name

        return {
            "serialNo": serial_no,
            "category": category_name,
            "subcategory": subcategory_name,
            "description": description,
            "spec": spec,
            "quantity": quantity or 1,
            "unitPrice": unit_price,
            "lineTotal": total,
        }

    for raw_line in text_lines:
        line = normalize_text(raw_line)
        if not line:
            continue

        lower = normalize_key(line)
        if not started:
            if "s/n" in lower and "category" in lower and "description" in lower:
                started = True
                quote = extract_quote_meta_from_intro("\n".join(intro_lines), "")
            else:
                intro_lines.append(line)
            continue

        is_category_header = lower in {normalize_key(name) for name in category_names_sorted}
        is_subtotal_line = "subtotal" in lower

        if "terms and conditions" in lower:
            if pending:
                row = finalize_pending(pending)
                if row:
                    rows.append(row)
            break

        if is_category_header or is_subtotal_line:
            if pending:
                row = finalize_pending(pending)
                if row:
                    rows.append(row)
                pending = ""
            if is_category_header:
                current_category = category_map[lower]
            continue

        if re.fullmatch(r"\d+\s+.*", line):
            if pending:
                row = finalize_pending(pending)
                if row:
                    rows.append(row)
            pending = line
            continue

        if pending:
            pending = f"{pending} {line}"
        else:
            pending = line

    if pending:
        row = finalize_pending(pending)
        if row:
            rows.append(row)

    return quote, rows


def parse_pdf_import(file_stream, library_payload=None):
    extracted_rows = []
    quote = {
        "preparedBy": "",
        "preparedPhone": "",
        "preparedEmail": "",
        "to": "",
        "proposalFor": "",
        "city": "",
        "quoteDate": "",
    }

    if pdfplumber is not None:
        file_stream.seek(0)
        with pdfplumber.open(file_stream) as pdf:
            full_text = []
            for page in pdf.pages:
                full_text.append(page.extract_text() or "")
                for table in page.extract_tables() or []:
                    if not table or not table[0]:
                        continue
                    header = [normalize_text(cell) for cell in table[0]]
                    if header[:3] == ["S/N", "Category", "Subcategory"]:
                        for row in table[1:]:
                            values = [normalize_text(cell) for cell in row]
                            if len(values) < 8 or not any(values):
                                continue
                            if re.fullmatch(r"\d+", values[0]):
                                quantity = parse_amount(values[5])
                                unit_price = parse_amount(values[6])
                                extracted_rows.append(
                                    {
                                        "serialNo": int(values[0]),
                                        "category": values[1] or "GENERAL",
                                        "subcategory": values[2] or "General Items",
                                        "description": values[3],
                                        "spec": values[4],
                                        "quantity": quantity or 1,
                                        "unitPrice": unit_price,
                                        "lineTotal": parse_amount(values[7]) or (quantity * unit_price),
                                    }
                                )

            text_blob = "\n".join(full_text)
            quote = extract_quote_meta_from_intro(text_blob, "")
            if not extracted_rows:
                quote, extracted_rows = parse_pdf_text_rows(text_blob.splitlines(), library_payload)

    if not extracted_rows and PdfReader is not None:
        file_stream.seek(0)
        reader = PdfReader(file_stream)
        lines = []
        for page in reader.pages:
            lines.extend((page.extract_text() or "").splitlines())
        joined = "\n".join(lines)
        quote = extract_quote_meta_from_intro(joined, "")
        quote, extracted_rows = parse_pdf_text_rows(lines, library_payload)

    if not extracted_rows:
        raise ValueError("No editable quotation rows were detected in this PDF. If this is one of your saved quotation PDFs, send it once and I can tune the parser further.")

    return {"importedFrom": "PDF", "quote": quote, "lines": dedupe_import_lines(extracted_rows), "summary": {}}


def parse_json_import(raw_bytes):
    payload = json.loads(raw_bytes.decode("utf-8"))
    if not isinstance(payload, dict) or not payload.get("lines"):
        raise ValueError("This JSON file does not contain editable quotation data.")
    return {
        "importedFrom": "JSON",
        "quote": payload.get("quote", {}),
        "lines": dedupe_import_lines(payload.get("lines", [])),
        "summary": payload.get("summary", {}),
    }


def build_docx(payload):
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX export.")

    company = payload.get("company") or SEED_LIBRARY["company"]
    quote = payload.get("quote") or {}
    lines = payload.get("lines") or []
    summary = payload.get("summary") or {}
    logo_stream = decode_data_url_image(payload.get("logoData"))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(7.0))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Inches(2.0)
    header_table.columns[1].width = Inches(4.8)

    logo_cell = header_table.cell(0, 0)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_path = os.path.join(BASE_DIR, "logo.png")
    if logo_stream is not None:
        run = logo_cell.paragraphs[0].add_run()
        run.add_picture(logo_stream, width=Inches(1.8))
    elif os.path.exists(logo_path):
        run = logo_cell.paragraphs[0].add_run()
        run.add_picture(logo_path, width=Inches(1.6))
    else:
        apply_cell_text(logo_cell, company["name"], bold=True, size=12)

    title_cell = header_table.cell(0, 1)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = title_para.add_run(company["name"] + "\n")
    run.bold = True
    run.font.size = Pt(18)
    run = title_para.add_run(company.get("quote_title", "QUOTE"))
    run.bold = True
    run.font.size = Pt(24)

    intro = doc.add_table(rows=1, cols=2)
    intro.alignment = WD_TABLE_ALIGNMENT.CENTER
    intro.columns[0].width = Inches(4.75)
    intro.columns[1].width = Inches(2.2)

    left = intro.cell(0, 0)
    left.text = ""
    for line in [
        company.get("tagline", ""),
        quote.get("preparedBy", ""),
        quote.get("preparedPhone", ""),
        quote.get("preparedEmail", ""),
        "",
        f"TO       {quote.get('to', '')}",
        quote.get("proposalFor", ""),
        quote.get("city", ""),
    ]:
        p = left.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        if line.startswith("TO"):
            run.bold = True

    right = intro.cell(0, 1)
    apply_cell_text(right, quote.get("quoteDate") or datetime.date.today().strftime("%d-%B-%Y"), bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.55, 0.8, 0.85, 2.1, 0.7, 0.45, 0.8, 0.95]
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    set_table_borders(table, color="AEB7C1")

    headers = ["S/N", "Category", "Subcategory", "Description", "Spec / Unit", "Qty", "Unit Price", "Total"]
    hdr = table.rows[0].cells
    for idx, label in enumerate(headers):
        apply_cell_text(hdr[idx], label, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr[idx], "DCE6F1")

    grouped = OrderedDict()
    for line in lines:
        key = line.get("category", "GENERAL")
        grouped.setdefault(key, []).append(line)

    for category_name, category_lines in grouped.items():
        cat_row = table.add_row().cells
        cat_row[0].merge(cat_row[7])
        apply_cell_text(cat_row[0], category_name, bold=True, size=10)
        set_cell_shading(cat_row[0], "F6E7C1")

        for line in category_lines:
            row = table.add_row().cells
            apply_cell_text(row[0], line.get("serialNo", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
            apply_cell_text(row[1], line.get("category", ""))
            apply_cell_text(row[2], line.get("subcategory", ""))
            apply_cell_text(row[3], line.get("description", ""))
            apply_cell_text(row[4], line.get("spec", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
            apply_cell_text(row[5], line.get("quantity", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
            apply_cell_text(row[6], money(line.get("unitPrice", 0)), align=WD_ALIGN_PARAGRAPH.RIGHT)
            apply_cell_text(row[7], money(line.get("lineTotal", 0)), align=WD_ALIGN_PARAGRAPH.RIGHT)

        subtotal_row = table.add_row().cells
        subtotal_row[0].merge(subtotal_row[6])
        apply_cell_text(subtotal_row[0], f"{category_name} Subtotal", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        apply_cell_text(subtotal_row[7], money(sum(float(item.get("lineTotal", 0)) for item in category_lines)), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_shading(subtotal_row[0], "F4F6F8")
        set_cell_shading(subtotal_row[7], "F4F6F8")

    doc.add_paragraph("")

    totals = doc.add_table(rows=4, cols=2)
    totals.alignment = WD_TABLE_ALIGNMENT.RIGHT
    totals.columns[0].width = Inches(1.8)
    totals.columns[1].width = Inches(1.4)
    set_table_borders(totals, color="AEB7C1")
    total_rows = [
        ("Subtotal", summary.get("subtotal", 0)),
        ("Discount", summary.get("discountAmount", 0)),
        ("Tax", summary.get("taxAmount", 0)),
        ("Grand Total", summary.get("grandTotal", 0)),
    ]
    for idx, (label, value) in enumerate(total_rows):
        apply_cell_text(totals.cell(idx, 0), label, bold=(idx == 3), align=WD_ALIGN_PARAGRAPH.RIGHT)
        apply_cell_text(totals.cell(idx, 1), money(value), bold=(idx == 3), align=WD_ALIGN_PARAGRAPH.RIGHT)
        if idx == 3:
            set_cell_shading(totals.cell(idx, 0), "DCE6F1")
            set_cell_shading(totals.cell(idx, 1), "DCE6F1")

    doc.add_paragraph("")
    terms_heading = doc.add_paragraph()
    terms_heading.add_run("Terms and Conditions:").bold = True
    for index, term in enumerate(company.get("terms", []), start=1):
        doc.add_paragraph(f"{index}- {term}")

    doc.add_paragraph("")
    for line in company.get("closing", []):
        doc.add_paragraph(line)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(company.get("footer_line_1", "") + "\n")
    footer.add_run(company.get("footer_line_2", ""))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


@app.route("/")
def index():
    return render_template("index.html", seed_data=SEED_LIBRARY)


@app.route("/api/import/file", methods=["POST"])
def import_file():
    uploaded = request.files.get("file")
    if uploaded is None or not uploaded.filename:
        return jsonify({"error": "Choose a PDF, DOCX, or JSON file to import."}), 400

    extension = os.path.splitext(uploaded.filename.lower())[1]
    file_bytes = uploaded.read()
    library_payload = None
    if request.form.get("library"):
        try:
            library_payload = json.loads(request.form.get("library"))
        except json.JSONDecodeError:
            library_payload = None

    try:
        if extension == ".docx":
            result = parse_docx_import(io.BytesIO(file_bytes))
        elif extension == ".pdf":
            result = parse_pdf_import(io.BytesIO(file_bytes), library_payload)
        elif extension == ".json":
            result = parse_json_import(file_bytes)
        else:
            return jsonify({"error": "Unsupported file type. Import PDF, DOCX, or JSON."}), 400
    except Exception as exc:
        return jsonify({"error": str(exc)}), 400

    return jsonify(result)


@app.route("/api/export/docx", methods=["POST"])
def export_docx():
    payload = request.get_json() or {}
    lines = payload.get("lines") or []
    if not lines:
        return jsonify({"error": "Add at least one line item before exporting DOCX."}), 400

    try:
        file_stream = build_docx(payload)
    except Exception as exc:  # pragma: no cover
        return jsonify({"error": str(exc)}), 500

    quote = payload.get("quote") or {}
    safe_to = slugify(quote.get("to", "client"))
    date_part = datetime.date.today().strftime("%Y%m%d")
    filename = f"Reliance-Quote-{safe_to}-{date_part}.docx"
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/api/export/json", methods=["POST"])
def export_json():
    payload = request.get_json() or {}
    lines = payload.get("lines") or []
    if not lines:
        return jsonify({"error": "Add at least one line item before exporting JSON."}), 400

    file_bytes = json.dumps(payload, indent=2).encode("utf-8")
    quote = payload.get("quote") or {}
    safe_to = slugify(quote.get("to", "client"))
    date_part = datetime.date.today().strftime("%Y%m%d")
    filename = f"Reliance-Quote-{safe_to}-{date_part}.json"
    return send_file(
        io.BytesIO(file_bytes),
        as_attachment=True,
        download_name=filename,
        mimetype="application/json",
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
